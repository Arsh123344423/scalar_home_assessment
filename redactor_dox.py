from docx import Document
from redactor import replacement_map

def redact_docx(input_file, output_file):
    doc = Document(input_file)

    # Paragraphs
    for para in doc.paragraphs:

        text = para.text

        for old, new in replacement_map.items():
            if old == "Email":
                continue
            text = text.replace(old, new)

        if para.runs:
            para.runs[0].text = text

            for run in para.runs[1:]:
                run.text = ""

    # Tables
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    text = para.text

                    for old, new in replacement_map.items():
                        if old == "Email":
                            continue
                        text = text.replace(old, new)

                    if para.runs:
                        para.runs[0].text = text

                        for run in para.runs[1:]:
                            run.text = ""

    doc.save(output_file)